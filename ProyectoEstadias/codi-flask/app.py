from flask import Flask, render_template, request, jsonify, session, redirect, url_for, flash, send_file
from conexion import get_connection
import bcrypt
from werkzeug.security import generate_password_hash, check_password_hash
from datetime import datetime
import sqlite3
import os
from functools import wraps
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
import logging
from pathlib import Path
from docx.text.run import Run
from werkzeug.utils import secure_filename


app = Flask(__name__)
app.secret_key = 'CoDI'  

ADMIN_EMAIL = "oscardanielramireztorrez@gmail.com"
ADMIN_PASSWORD_HASH = generate_password_hash("123456")

def role_required(*roles):
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            if 'rol' not in session:
                return redirect(url_for('login'))
            if session['rol'] not in roles:
                return redirect(url_for('inicio'))
            return f(*args, **kwargs)
        return decorated_function
    return decorator 

@app.route('/', methods=['GET'])
def root():
    return redirect(url_for('login'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    if 'IdUsuario' in session:
        if session.get('rol') == 'administrador':
            return redirect(url_for('administrador'))
        elif session.get('rol') == 'jefe':
            return redirect(url_for('jefe'))
        elif session.get('rol') == 'recepcion':
            return redirect(url_for('recepcion'))
        return redirect(url_for('inicio'))

    if request.method == 'POST':
        correo = request.form.get('correo')
        contrasena = request.form.get('contrasena')

        if not correo or not contrasena:
            return render_template('login.html', error='Por favor, complete todos los campos.')

        # Verificación para admin
        if correo == ADMIN_EMAIL:
            if check_password_hash(ADMIN_PASSWORD_HASH, contrasena):
                session.clear()
                session['rol'] = 'administrador'
                session['correo'] = ADMIN_EMAIL
                return redirect(url_for('administrador'))
            return render_template('login.html', error='Credenciales inválidas')

        # Verificación para usuarios normales
        conn = get_connection()
        cursor = conn.cursor(dictionary=True)
        try:
            cursor.execute("""
                SELECT IdUsuario, nombre, correo, contrasena, cargo, siglas, area, rol 
                FROM usuarios 
                WHERE correo = %s
            """, (correo,))
            user = cursor.fetchone()

            if user:
                # Depuración: Verificar contraseña y rol
                app.logger.debug(f"Usuario encontrado: {user['nombre']}, Rol: {user['rol']}")
                
                if bcrypt.checkpw(contrasena.encode('utf-8'), user['contrasena'].encode('utf-8')):
                    session.clear()
                    session['IdUsuario'] = user['IdUsuario']
                    session['nombre'] = user['nombre']
                    session['correo'] = user['correo']
                    session['cargo'] = user['cargo']
                    session['siglas'] = user['siglas']
                    session['area'] = user['area']
                    session['rol'] = user['rol']

                    app.logger.debug(f"Login exitoso. Redirigiendo según rol: {user['rol']}")
                    
                    if user['rol'] == 'administrador':
                        return redirect(url_for('administrador'))
                    elif user['rol'] == 'jefe':
                        return redirect(url_for('jefe'))
                    elif user['rol'] == 'recepcion':
                        return redirect(url_for('recepcion'))
                    elif user['rol'] == 'usuario' :
                        return redirect(url_for('inicio'))
                    return redirect(url_for('inicio'))
            
            return render_template('login.html', error='Credenciales inválidas')
        except Exception as e:
            app.logger.error(f"Error en login: {str(e)}")
            return render_template('login.html', error='Error al iniciar sesión')
        finally:
            cursor.close()
            conn.close()

    return render_template('login.html')


@app.route('/logout')
def logout():
    session.clear() 
    return redirect(url_for('login'))

@app.route('/index')
def inicio():
    if 'IdUsuario' not in session and 'rol' not in session:
        return redirect(url_for('login'))

    if 'IdUsuario' in session:
        conn = get_connection()
        cursor = conn.cursor(dictionary=True)
        try:
            cursor.execute("""
                SELECT nombre, correo, cargo, siglas, area
                FROM usuarios
                WHERE IdUsuario = %s
            """, (session['IdUsuario'],))
            user_data = cursor.fetchone()
            
            if not user_data:
                return redirect(url_for('logout'))

            session['nombre'] = user_data['nombre']
            session['correo'] = user_data['correo']
            session['cargo'] = user_data['cargo']
            session['siglas'] = user_data['siglas']
            session['area'] = user_data['area']

            return render_template('index.html')
            
        except Exception as e:
            app.logger.error(f"Error al cargar index: {str(e)}")
            return redirect(url_for('logout'))
        finally:
            cursor.close()
            conn.close()
    
    return redirect(url_for('administrador'))

@app.route('/formatos')
def formatos():
    if 'IdUsuario' not in session or 'rol' not in session:
        return redirect(url_for('login'))
    return render_template('formatos.html')


@app.route('/copiar_doc', methods=['GET'])
def copiar_doc():
    try:
        template_path = "static/documentos/Formato de oficio.docx"
        if not os.path.exists(template_path):
            return jsonify({'error': 'El documento no existe'}), 404

        doc = Document(template_path)

        contenido = []
        for p in doc.paragraphs:
            if p.text.strip():
                contenido.append(p.text.strip())

        texto = "\n".join(contenido)

        return jsonify({"texto": texto})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/subir_documento', methods=['POST'])
def subir_documento():
    if 'IdUsuario' not in session or 'rol' not in session:
        return redirect(url_for('login'))

    if 'documento' not in request.files:
        return jsonify({'error': 'No se envió archivo'}), 400

    archivo = request.files['documento']

    if archivo.filename == '':
        return jsonify({'error': 'Archivo no válido'}), 400

    # Ruta fija donde se guarda SIEMPRE
    destino = os.path.join("static", "documentos", "Formato de oficio.docx")

    try:
        archivo.save(destino)
        return jsonify({'mensaje': 'Formato reemplazado correctamente'}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500



@app.route('/admin', methods=['GET'])
def administrador():
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM usuarios")
    usuarios = cursor.fetchall()
    cursor.close()
    conn.close()
    return render_template('admin.html', usuarios=usuarios)

@app.route('/jefe')
@role_required('jefe', 'administrador')  # Permite ambos roles
def jefe():
    conn = get_connection()
    cursor = conn.cursor(dictionary=True)
    
    try:
        cursor.execute("""
            SELECT nombre, correo, cargo, siglas, area
            FROM usuarios
            WHERE IdUsuario = %s
        """, (session['IdUsuario'],))
        user_data = cursor.fetchone()
        
        if not user_data:
            return redirect(url_for('logout'))

        return render_template('jefe.html', user=user_data)
        
    except Exception as e:
        app.logger.error(f"Error en panel jefe: {str(e)}")
        return redirect(url_for('inicio'))
    finally:
        cursor.close()
        conn.close()

@app.route('/recepcion')
@role_required('recepcion', 'administrador')  # Permite ambos roles
def recepcion():
    conn = get_connection()
    cursor = conn.cursor(dictionary=True)
    
    try:
        cursor.execute("""
            SELECT nombre, correo, cargo, siglas, area
            FROM usuarios
            WHERE IdUsuario = %s
        """, (session['IdUsuario'],))
        user_data = cursor.fetchone()
        
        if not user_data:
            return redirect(url_for('logout'))

        return render_template('recepcion.html', user=user_data)
        
    except Exception as e:
        app.logger.error(f"Error en panel recepción: {str(e)}")
        return redirect(url_for('inicio'))
    finally:
        cursor.close()
        conn.close()


@app.route('/recepcion/guardar_entrada', methods=['POST'])
@role_required('recepcion', 'administrador')
def recepcion_guardar_entrada():
    if 'IdUsuario' not in session:
        return redirect(url_for('login'))

    numero = request.form.get('numero')
    area_origen = request.form.get('area')
    resumen = request.form.get('descripcion')
    
    if not numero or not area_origen or not resumen:
        return jsonify({"success": False, "message": "Faltan datos"}), 400

    conn = get_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("""
            INSERT INTO entrada (numero, area_origen, resumen)
            VALUES (%s, %s, %s)
        """, (numero, area_origen, resumen))
        conn.commit()
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "message": str(e)}), 500
    finally:
        cursor.close()
        conn.close()

    return jsonify({"success": True, "message": "Entrada guardada correctamente"}), 201

@app.route('/api/entradas', methods=['GET'])
@role_required('recepcion', 'administrador')
def listar_entradas():
    conn = get_connection()
    cursor = conn.cursor(dictionary=True)
    try:
        cursor.execute("SELECT * FROM entrada ORDER BY IdEntrada DESC")
        entradas = cursor.fetchall()
        return jsonify(entradas)
    finally:
        cursor.close()
        conn.close()


@app.route("/api/documentos", methods=["GET"])
def api_listar_documentos():
    conn = get_connection()
    cursor = conn.cursor(dictionary=True)
    cursor.execute("SELECT * FROM documentos_envio ORDER BY id DESC")
    documentos = cursor.fetchall()
    cursor.close()
    conn.close()
    return jsonify(documentos)

@app.route("/api/documentos", methods=["POST"])
def api_guardar_documento():
    numero = request.form["numero"]
    area = request.form["area"]
    descripcion = request.form["descripcion"]
    fecha = request.form["fecha"]
  
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        "INSERT INTO documentos_envio (numero, area, descripcion, fecha) VALUES (%s, %s, %s, %s)",
        (numero, area, descripcion, fecha)
    )
    conn.commit()
    cursor.close()
    conn.close()
    return jsonify({"message": "Documento guardado con éxito"}), 201


@app.route('/guardar_documento', methods=['POST'])
def guardar_documento():
    if 'IdUsuario' not in session:
        return jsonify({'success': False, 'message': 'No autorizado'}), 401

    data = request.get_json()
    numero = data.get('numero')
    area = data.get('area')
    descripcion = data.get('descripcion')
    
    if not numero or not area or not descripcion:
        return jsonify({'success': False, 'message': 'Faltan datos'}), 400
    
    fecha = datetime.now().strftime('%Y-%m-%d')

    conn = get_connection()
    cursor = conn.cursor()
    try:
        cursor.execute(
            "INSERT INTO documentos_envio (numero, area, descripcion, fecha) VALUES (%s, %s, %s, %s)",
            (numero, area, descripcion, fecha)
        )
        conn.commit()
    except Exception as e:
        conn.rollback()
        return jsonify({'success': False, 'message': str(e)}), 500
    finally:
        cursor.close()
        conn.close()

    return jsonify({'success': True, 'message': 'Documento guardado correctamente'})
@app.route('/Entrada')
def ver_documentos():
    conn = get_connection()
    cursor = conn.cursor(dictionary=True)
    cursor.execute("SELECT * FROM documentos_envio ORDER BY id DESC")
    documentos = cursor.fetchall()
    cursor.close()
    conn.close()
    return jsonify(documentos)





@app.route('/documento')
def editar_documento():
    try:
        folio_id = request.args.get('folio_id')
        if not folio_id:
            flash('No se ha especificado un folio', 'error')
            return redirect(url_for('folios'))

        conn = get_connection()
        cursor = conn.cursor(dictionary=True)

        # Datos del folio
        cursor.execute("""
            SELECT s.*, u.nombre as usuario_nombre, u.IdArea
            FROM salidas s
            JOIN usuarios u ON s.IdUsuario = u.IdUsuario
            WHERE s.IdSalida = %s
        """, (folio_id,))
        folio = cursor.fetchone()
        if not folio:
            flash('Folio no encontrado', 'error')
            return redirect(url_for('folios'))

        # Todas las áreas
        cursor.execute("SELECT IdArea as id, nombre, UI, CC FROM areas ORDER BY nombre")
        areas = cursor.fetchall()

        # Destinatarios disponibles
        cursor.execute("SELECT IdDestino as id, nombre, cargo FROM destinatarios ORDER BY nombre")
        destinatarios = cursor.fetchall()

        # Unidades disponibles
        cursor.execute("SELECT IdUnidad as id, nombre FROM unidades ORDER BY nombre")
        unidades = cursor.fetchall()

        cursor.close()
        conn.close()

        return render_template('documento.html',
                               folio=folio,
                               folio_id=folio_id,
                               areas=areas,
                               destinatarios=destinatarios,
                               unidades=unidades)

    except Exception as e:
        print(f"ERROR en /documento: {str(e)}")
        flash('Error al cargar el editor de documento', 'error')
        return redirect(url_for('folios'))



@app.route('/generar_documento', methods=['POST'])
def generar_documento():
    conn = None
    cursor = None
    try:
        if 'IdUsuario' not in session:
            return jsonify({'error': 'No autorizado'}), 401

        if not request.is_json:
            return jsonify({'error': 'Se esperaba formato JSON'}), 400

        data = request.get_json()

        # Validar campos obligatorios
        required_fields = ['folio_id', 'destinatarios', 'asunto', 'cuerpo']
        for campo in required_fields:
            if campo not in data or not data[campo]:
                return jsonify({'error': f'Campo {campo} es requerido'}), 400

        if not data['destinatarios']:
            return jsonify({'error': 'Se requiere al menos un destinatario'}), 400

        conn = get_connection()
        cursor = conn.cursor(dictionary=True)

        # Obtener folio
        cursor.execute("SELECT numero FROM salidas WHERE IdSalida = %s", (data['folio_id'],))
        folio_data = cursor.fetchone()
        if not folio_data:
            return jsonify({'error': 'Folio no encontrado'}), 404

        numero_folio = folio_data['numero']
        partes = numero_folio.split('/')
        siglas = partes[0] if len(partes) > 0 else "CI"
        consecutivo = partes[1] if len(partes) > 1 else str(data['folio_id']).zfill(4)
        anio = partes[2] if len(partes) > 2 else str(datetime.now().year)

        # Tomar UI y CC enviados desde el front-end
        ui = data.get('ui', '129001')
        cc = data.get('cc', '500100')
        numero_oficio = f"{ui}/{cc}/{siglas}/{consecutivo}/{anio}"

        # Fecha formateada
        meses = [
            'enero', 'febrero', 'marzo', 'abril', 'mayo', 'junio',
            'julio', 'agosto', 'septiembre', 'octubre', 'noviembre', 'diciembre'
        ]
        now = datetime.now()
        mes_nombre = meses[now.month-1] if 1 <= now.month <= 12 else 'enero'
        fecha_formateada = f"Acapulco de Juárez, Gro., a {now.day} de {mes_nombre} de {now.year}"

        # Cargar plantilla
        template_path = "static/documentos/Formato de oficio.docx"
        if not os.path.exists(template_path):
            return jsonify({'error': 'Plantilla de documento no encontrada'}), 500

        doc = Document(template_path)

        # Preparar destinatarios
        destinatarios_texto = ""
        for dest in data['destinatarios']:
            nombre = dest.get('nombre', '')
            cargo = dest.get('cargo', '')
            destinatarios_texto += f"{nombre}\n{cargo}\n"
        destinatarios_texto += ""


        # Preparar reemplazos
        replacements = {
            "{{NUMERO_OFICIO}}": f"Of. N° {numero_oficio}",
            "{{FECHA}}": fecha_formateada,
            "{{UNIDAD}}": data.get('unidad_nombre', ''),
            "{{AREA}}": data.get('area_nombre', ''),
            "{{DESTINATARIOS}}": destinatarios_texto.strip(),
            "{{ASUNTO}}": data.get('asunto', ''),
            "{{CUERPO}}": data.get('cuerpo', ''),
            "{{ELABORADOR}}": data.get('elaborador', '')
        }

        # Función para reemplazar texto en todo el documento
        def replace_text_in_doc(doc, replacements):
            for paragraph in doc.paragraphs:
                for search, replace in replacements.items():
                    if search in paragraph.text:
                        full_text = paragraph.text.replace(search, replace)
                        for run in paragraph.runs:
                            run.text = ''
                        if paragraph.runs:
                            paragraph.runs[0].text = full_text              
                            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for search, replace in replacements.items():
                                if search in full_text:
                                    full_text = full_text.replace(search, replace)
                                    for run in paragraph.runs:
                                        run.text = ''
                                    if paragraph.runs:
                                        paragraph.runs[0].text = full_text

        # Aplicar reemplazos
        replace_text_in_doc(doc, replacements)

        # Guardar en memoria y enviar al cliente
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        nombre_archivo = f"Oficio_{numero_oficio.replace('/', '_')}.docx"

        return send_file(
            file_stream,
            as_attachment=True,
            download_name=nombre_archivo,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        logging.error(f'Error al generar documento: {str(e)}', exc_info=True)
        return jsonify({'error': f'Error al generar el documento: {str(e)}'}), 500
    finally:
        if cursor:
            cursor.close()
        if conn and conn.is_connected():
            conn.close()








              
              
              
              
              
              
              
              
              
              
              
              
              
              
@app.route('/add_user', methods=['POST'])
def add_user():
    data = request.get_json()
    conn = get_connection()
    cursor = conn.cursor()

    try:
        hashed_password = bcrypt.hashpw(data['contrasena'].encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

        cursor.execute("""
            INSERT INTO usuarios (nombre, correo, contrasena, cargo, siglas, area, rol)
            VALUES (%s, %s, %s, %s, %s, %s, %s)
        """, (data['nombre'], data['correo'], hashed_password, data['cargo'], data['siglas'], data['area'], data['rol']))
        
        new_id = cursor.lastrowid 
        conn.commit()
        return jsonify({"message": "Usuario agregado exitosamente.", "id": new_id}), 201
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    finally:
        cursor.close()
        conn.close()

@app.route('/edit_user/<int:id>', methods=['POST'])
def edit_user(id):
    conn = get_connection()
    cursor = conn.cursor()

    try:
        data = request.get_json()
        if 'contrasena' in data and data['contrasena']:
            hashed_password = bcrypt.hashpw(data['contrasena'].encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
            cursor.execute("""
                UPDATE usuarios
                SET nombre = %s, correo = %s, contrasena = %s, cargo = %s, siglas = %s, area = %s, rol = %s
                WHERE Idusuario = %s
            """, (data['nombre'], data['correo'], hashed_password, data['cargo'], data['siglas'], data['area'], data['rol'], id))
        else:
            cursor.execute("""
                UPDATE usuarios
                SET nombre = %s, correo = %s, cargo = %s, siglas = %s, area = %s, rol = %s
                WHERE Idusuario = %s
            """, (data['nombre'], data['correo'], data['cargo'], data['siglas'], data['area'], data['rol'], id))

        conn.commit()
        return jsonify({"message": "Usuario actualizado exitosamente."}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    finally:
        cursor.close()
        conn.close()

@app.route('/delete_user/<int:id>', methods=['POST'])
def delete_user(id):
       print(f"Eliminando usuario con ID: {id}")  
       conn = get_connection()
       cursor = conn.cursor()

       try:
           cursor.execute("DELETE FROM usuarios WHERE Idusuario = %s", (id,))
           conn.commit()
           return jsonify({"message": "Usuario eliminado exitosamente."}), 204
       except Exception as e:
           print(f"Error al eliminar: {str(e)}")  
           return jsonify({"error": str(e)}), 400
       finally:
           cursor.close()
           conn.close()
   
@app.route('/add_unidad', methods=['POST'])
def add_unidad():
    data = request.get_json()
    conn = get_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("""
            INSERT INTO unidades (nombre, descripcion, localidad)
            VALUES (%s, %s, %s)
        """, (data['nombre'], data['descripcion'], data['localidad']))
        conn.commit()
        
        return jsonify({"message": "Unidad agregada exitosamente."}), 201
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    finally:
        cursor.close()
        conn.close()

@app.route('/edit_unidad/<int:id>', methods=['POST'])
def edit_unidad(id):
    data = request.get_json()
    conn = get_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("""
            UPDATE unidades
            SET nombre = %s, descripcion = %s, localidad = %s
            WHERE Idunidad = %s
        """, (data['nombre'], data['descripcion'], data['localidad'], id))
        conn.commit()
        
        return jsonify({"message": "Unidad actualizada exitosamente."}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    finally:
        cursor.close()
        conn.close()

@app.route('/delete_unidad/<int:id>', methods=['POST'])
def delete_unidad(id):
    conn = get_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("DELETE FROM unidades WHERE Idunidad = %s", (id,))
        conn.commit()
        return jsonify({"message": "Unidad eliminada exitosamente."}), 204
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    finally:
        cursor.close()
        conn.close()

@app.route('/get_unidades', methods=['GET'])
def get_unidades():
    conn = get_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("SELECT Idunidad, nombre, descripcion, localidad, fecha FROM unidades")
        unidades = cursor.fetchall()
        
        result = []
        for unidad in unidades:
            result.append({
                "id": unidad[0],
                "nombre": unidad[1],
                "descripcion": unidad[2],
                "localidad": unidad[3],
                "fecha": unidad[4].isoformat() 
            })
        
        return jsonify(result), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    finally:
        cursor.close()
        conn.close()
        
@app.route('/get_areas', methods=['GET'])
def get_areas():
    conn = get_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("SELECT IdArea, nombre, UI, CC, siglas, formato, fecha FROM areas")
        areas = cursor.fetchall()
        
        result = []
        for area in areas:
            result.append({
                "id": area[0],
                "nombre": area[1],
                "UI": area[2],
                "CC": area[3],
                "siglas": area[4],
                "formato": area[5],
                "fecha": area[6]
            })
        
        return jsonify(result), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    finally:
        cursor.close()
        conn.close()

@app.route('/add_area', methods=['POST'])
def add_area():
    data = request.get_json()
    conn = get_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("""
            INSERT INTO areas (nombre, UI, CC, siglas, formato, fecha)
            VALUES (%s, %s, %s, %s, %s, %s)
        """, (data['nombre'], data['UI'], data['CC'], data['siglas'], data['formato'], data['fecha']))
        conn.commit()
        
        return jsonify({"message": "Área agregada exitosamente."}), 201
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    finally:
        cursor.close()
        conn.close()

@app.route('/edit_area/<int:id>', methods=['POST'])
def edit_area(id):
    data = request.get_json()
    conn = get_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("""
            UPDATE areas
            SET nombre = %s, UI = %s, CC = %s, siglas = %s, formato = %s, fecha = %s
            WHERE IdArea = %s
        """, (data['nombre'], data['UI'], data['CC'], data['siglas'], data['formato'], data['fecha'], id))
        conn.commit()
        
        return jsonify({"message": "Área actualizada exitosamente."}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    finally:
        cursor.close()
        conn.close()

@app.route('/delete_area/<int:id>', methods=['POST'])
def delete_area(id):
    conn = get_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("DELETE FROM areas WHERE IdArea = %s", (id,))
        conn.commit()
        return jsonify({"message": "Área eliminada exitosamente."}), 204
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    finally:
        cursor.close()
        conn.close()

@app.route('/destinatarios')
def mostrar_destinatarios():
    return render_template('destinatarios.html')

@app.route('/api/destinatarios', methods=['GET', 'POST'])
def manejar_destinatarios():
    conn = get_connection()
    if not conn:
        return jsonify({'error': 'Error de conexión a la base de datos'}), 500

    try:
        cursor = conn.cursor(dictionary=True)
        
        if request.method == 'GET':
            page = request.args.get('page', 1, type=int)
            per_page = 5
            search = request.args.get('search', '', type=str)
            
            query = 'SELECT IdDestino AS id, nombre, cargo FROM destinatarios'
            params = []
            
            if search:
                query += ' WHERE nombre LIKE %s OR cargo LIKE %s'
                params.extend([f'%{search}%', f'%{search}%'])
            
            query += ' LIMIT %s OFFSET %s'
            params.extend([per_page, (page - 1) * per_page])
            
            cursor.execute(query, params)
            destinatarios = cursor.fetchall()
            
            count_query = 'SELECT COUNT(*) AS total FROM destinatarios'
            if search:
                count_query += ' WHERE nombre LIKE %s OR cargo LIKE %s'
                cursor.execute(count_query, [f'%{search}%', f'%{search}%'])
            else:
                cursor.execute(count_query)
                
            total_count = cursor.fetchone()['total']
            total_pages = (total_count + per_page - 1) // per_page
            
            return jsonify({
                'destinatarios': destinatarios,
                'total_pages': total_pages,
                'current_page': page
            })
            
        elif request.method == 'POST':
            data = request.get_json()
            nombre = data.get('nombre')
            cargo = data.get('cargo')
            
            if not nombre or not cargo:
                return jsonify({'error': 'Nombre y cargo son requeridos'}), 400
            
            cursor.execute(
                'INSERT INTO destinatarios (nombre, cargo) VALUES (%s, %s)',
                (nombre, cargo)
            )
            destinatario_id = cursor.lastrowid
            conn.commit()
            
            cursor.execute(
                'SELECT IdDestino AS id, nombre, cargo FROM destinatarios WHERE IdDestino = %s', 
                (destinatario_id,)
            )
            destinatario = cursor.fetchone()
            
            return jsonify(destinatario), 201
            
    except Exception as e:
        conn.rollback()
        return jsonify({'error': str(e)}), 500
    finally:
        cursor.close()
        conn.close()

@app.route('/api/destinatarios/<int:id>', methods=['PUT', 'DELETE'])
def manejar_destinatario(id):
    conn = get_connection()
    if not conn:
        return jsonify({'error': 'Error de conexión a la base de datos'}), 500

    try:
        cursor = conn.cursor(dictionary=True)
        
        if request.method == 'PUT':
            data = request.get_json()
            nombre = data.get('nombre')
            cargo = data.get('cargo')
            
            if not nombre or not cargo:
                return jsonify({'error': 'Nombre y cargo son requeridos'}), 400
            
            cursor.execute(
                'UPDATE destinatarios SET nombre = %s, cargo = %s WHERE IdDestino = %s',
                (nombre, cargo, id)
            )
            
            if cursor.rowcount == 0:
                return jsonify({'error': 'Destinatario no encontrado'}), 404
            
            conn.commit()
            
            cursor.execute(
                'SELECT IdDestino AS id, nombre, cargo FROM destinatarios WHERE IdDestino = %s', 
                (id,)
            )
            destinatario = cursor.fetchone()
            
            return jsonify(destinatario)
            
        elif request.method == 'DELETE':
            cursor.execute('DELETE FROM destinatarios WHERE IdDestino = %s', (id,))
            
            if cursor.rowcount == 0:
                return jsonify({'error': 'Destinatario no encontrado'}), 404
            
            conn.commit()
            return jsonify({'message': 'Destinatario eliminado correctamente'}), 200
            
    except Exception as e:
        conn.rollback()
        return jsonify({'error': str(e)}), 500
    finally:
        cursor.close()
        conn.close()
        
@app.route('/folios')
def folios():
    """Ruta para servir la página HTML"""
    return render_template('folios.html')

@app.route('/api/folios', methods=['GET', 'POST'])
def manejar_folios():
    conn = get_connection()
    if not conn:
        return jsonify({'error': 'Error de conexión a la base de datos'}), 500

    try:
        cursor = conn.cursor(dictionary=True)
        
        if request.method == 'GET':
            # Verificar autenticación
            if 'IdUsuario' not in session:
                return jsonify({'error': 'No autorizado'}), 401
                
            # Obtener información del usuario
            user_id = session['IdUsuario']
            user_role = session.get('rol', 'usuario')
            user_area = session.get('area')

            # Manejar paginación y búsqueda
            page = request.args.get('page', 1, type=int)
            per_page = 5
            search = request.args.get('search', '', type=str)
            
            # Construir la consulta base con JOIN para obtener el área del creador
            query = """
                SELECT s.IdSalida AS id, s.numero, s.dirigido, s.asunto, s.tipo, s.fecha_registro
                FROM salidas s
                JOIN usuarios u ON s.IdUsuario = u.IdUsuario
                WHERE 1=1
            """
            params = []
            
            # Aplicar filtros según el rol del usuario
            if user_role == 'usuario':
                query += ' AND s.IdUsuario = %s'
                params.append(user_id)
            elif user_role in ('recepcion', 'jefe'):
                query += ' AND u.area = %s'
                params.append(user_area)
            
            # Aplicar búsqueda si existe
            if search:
                query += ' AND (s.numero LIKE %s OR s.dirigido LIKE %s OR s.asunto LIKE %s)'
                params.extend([f'%{search}%', f'%{search}%', f'%{search}%'])
            
            # Ordenar y paginar
            query += ' ORDER BY s.fecha_registro DESC LIMIT %s OFFSET %s'
            params.extend([per_page, (page - 1) * per_page])
            
            cursor.execute(query, params)
            folios = cursor.fetchall()
            
            # Contar total de folios para paginación (con los mismos filtros)
            count_query = """
                SELECT COUNT(*) AS total 
                FROM salidas s
                JOIN usuarios u ON s.IdUsuario = u.IdUsuario
                WHERE 1=1
            """
            count_params = []
            
            if user_role == 'usuario':
                count_query += ' AND s.IdUsuario = %s'
                count_params.append(user_id)
            elif user_role in ('recepcion', 'jefe'):
                count_query += ' AND u.area = %s'
                count_params.append(user_area)
            
            if search:
                count_query += ' AND (s.numero LIKE %s OR s.dirigido LIKE %s OR s.asunto LIKE %s)'
                count_params.extend([f'%{search}%', f'%{search}%', f'%{search}%'])
                
            cursor.execute(count_query, count_params)
            total_count = cursor.fetchone()['total']
            total_pages = (total_count + per_page - 1) // per_page
            
            return jsonify({
                'folios': folios,
                'totalPages': total_pages
            })
            
        elif request.method == 'POST':
            # Validar que el usuario esté autenticado
            if 'IdUsuario' not in session:
                return jsonify({'error': 'No autorizado'}), 401
                
            data = request.get_json()
            user_id = session['IdUsuario']
            user_siglas = session.get('siglas', 'DS')
            
            # Validar campos requeridos
            required_fields = ['dirigido', 'asunto', 'tipo']
            if not all(field in data for field in required_fields):
                return jsonify({'error': 'Faltan campos requeridos'}), 400
                
            # Obtener el próximo número de folio
            year = datetime.now().year
            tipo = data['tipo']
            
            cursor.execute("""
                SELECT MAX(CAST(SUBSTRING_INDEX(SUBSTRING_INDEX(numero, '/', 2), '/', -1) AS UNSIGNED)) as max_num
                FROM salidas
                WHERE tipo = %s AND YEAR(fecha_registro) = %s 
                AND numero LIKE %s
            """, (tipo, year, f'{user_siglas}/%/{year}'))

            result = cursor.fetchone()
            next_num = (result['max_num'] or 0) + 1
            
            # Formatear número de folio
            numero_folio = f"{user_siglas}/{str(next_num).zfill(4)}/{year}"
            
            # Insertar el nuevo folio (incluyendo el IdUsuario)
            cursor.execute("""
                INSERT INTO salidas (
                    numero, dirigido, asunto, tipo, fecha_registro, IdUsuario
                ) VALUES (%s, %s, %s, %s, CURDATE(), %s)
            """, (
                numero_folio, 
                data['dirigido'], 
                data['asunto'], 
                tipo,
                user_id
            ))
            
            folio_id = cursor.lastrowid
            conn.commit()
            
            # Devolver el folio creado
            cursor.execute("""
                SELECT IdSalida AS id, numero, dirigido, asunto, tipo, fecha_registro
                FROM salidas 
                WHERE IdSalida = %s
            """, (folio_id,))
            folio = cursor.fetchone()
            
            return jsonify(folio), 201
            
    except Exception as e:
        conn.rollback()
        return jsonify({'error': str(e)}), 500
    finally:
        cursor.close()
        conn.close()

@app.route('/api/folios/<int:id>', methods=['PUT', 'DELETE'])
def manejar_folio(id):
    conn = get_connection()
    if not conn:
        return jsonify({'error': 'Error de conexión a la base de datos'}), 500

    try:
        cursor = conn.cursor(dictionary=True)
        
        # Validar que el usuario esté autenticado
        if 'IdUsuario' not in session:
            return jsonify({'error': 'No autorizado'}), 401
            
        user_id = session['IdUsuario']
        user_role = session.get('rol', 'usuario')
        user_area = session.get('area')
        
        # Verificar permisos antes de cualquier operación
        cursor.execute("""
            SELECT s.IdSalida, u.area, s.IdUsuario
            FROM salidas s
            JOIN usuarios u ON s.IdUsuario = u.IdUsuario
            WHERE s.IdSalida = %s
        """, (id,))
        folio = cursor.fetchone()
        
        if not folio:
            return jsonify({'error': 'Folio no encontrado'}), 404
            
        # Verificar permisos según rol
        if user_role == 'usuario':
            # Usuario solo puede modificar/eliminar sus propios folios
            if folio['IdUsuario'] != user_id:
                return jsonify({'error': 'No tienes permisos para esta acción'}), 403
        elif user_role in ('recepcion', 'jefe'):
            # Recepción y jefe solo pueden modificar folios de su área
            if folio['area'] != user_area:
                return jsonify({'error': 'No tienes permisos para esta acción'}), 403
        
        if request.method == 'PUT':
            data = request.get_json()
            
            # Validar campos requeridos
            required_fields = ['dirigido', 'asunto', 'tipo']
            if not all(field in data for field in required_fields):
                return jsonify({'error': 'Faltan campos requeridos'}), 400
                
            # Actualizar el folio
            cursor.execute("""
                UPDATE salidas 
                SET dirigido = %s, asunto = %s, tipo = %s
                WHERE IdSalida = %s
            """, (
                data['dirigido'], 
                data['asunto'], 
                data['tipo'],
                id
            ))
            
            conn.commit()
            
            # Devolver el folio actualizado
            cursor.execute("""
                SELECT IdSalida AS id, numero, dirigido, asunto, tipo, fecha_registro
                FROM salidas 
                WHERE IdSalida = %s
            """, (id,))
            folio = cursor.fetchone()
            
            return jsonify(folio)
            
        elif request.method == 'DELETE':
            cursor.execute('DELETE FROM salidas WHERE IdSalida = %s', (id,))
            conn.commit()
            return jsonify({'message': 'Folio eliminado correctamente'}), 200
            
    except Exception as e:
        conn.rollback()
        return jsonify({'error': str(e)}), 500
    finally:
        cursor.close()
        conn.close()

@app.route('/api/folios/next-number', methods=['GET'])
def obtener_proximo_numero():
    conn = get_connection()
    if not conn:
        return jsonify({'error': 'Error de conexión a la base de datos'}), 500

    try:
        cursor = conn.cursor(dictionary=True)
        
        # Validar que el usuario esté autenticado
        if 'IdUsuario' not in session:
            return jsonify({'error': 'No autorizado'}), 401
            
        # Obtener parámetros según el frontend
        tipo = request.args.get('tipo', 'interno')
        year = request.args.get('year', datetime.now().year, type=int)
        siglas = session.get('siglas', 'DS')  # Usamos las siglas del usuario logueado
        
        # Consulta modificada para el nuevo formato SIGLAS-NUMERO-AÑO
        cursor.execute("""
            SELECT MAX(CAST(SUBSTRING(numero, LOCATE('-', numero) + 1, 
            LOCATE('-', numero, LOCATE('-', numero) + 1) - LOCATE('-', numero) - 1) AS UNSIGNED)) as max_num 
            FROM salidas 
            WHERE tipo = %s AND YEAR(fecha_registro) = %s 
            AND numero LIKE %s
        """, (tipo, year, f'{siglas}-%-{year}'))
        result = cursor.fetchone()
        
        next_num = (result['max_num'] or 0) + 1
        
        # Formatear número de ejemplo según el frontend
        numero_ejemplo = f"{siglas}-{str(next_num).zfill(4)}/{year}"
        
        return jsonify({
            'nextNumber': next_num,
            'formattedNumber': numero_ejemplo
        })
        

    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        cursor.close()
        conn.close()
                
 
if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
    #app.run(debug=True, port=5000)
    #app.run(debug=True, host='192.168.1.151', port=5000, ssl_context=('192.168.1.151.pem', '192.168.1.151-key.pem'))
    #app.run(debug=True, host='192.168.1.92', port=5000, ssl_context=('192.168.1.92.pem', '192.168.1.92-key.pem'))